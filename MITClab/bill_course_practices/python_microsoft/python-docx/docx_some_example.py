from docx import Document
from docx.shared import Inches

document = Document()

# add a paragraph

paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')

document.add_heading('The REAL meaning of the universe')
document.add_heading('The role of dolphins', level=2)
document.add_page_break()

# ======================================================

# add a table
table = document.add_table(rows=2, cols=2)
cell = table.cell(0, 1)
cell.text = 'parrot, possibly dead'
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'
for row in table.rows:
    for cell in row.cells:
        print(cell.text)

row_count = len(table.rows)
col_count = len(table.columns)
row = table.add_row()

# get table data -------------
items = (
    (7, '1024', 'Plush kittens'),
    (3, '2042', 'Furbees'),
    (1, '1288', 'French Poodle Collars, Deluxe'),
)

# add table ------------------
table = document.add_table(1, 3)

# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Qty'
heading_cells[1].text = 'SKU'
heading_cells[2].text = 'Description'

# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = str(item.qty)
    cells[1].text = item.sku
    cells[2].text = item.desc

table.style = 'LightShading-Accent1'

# ======================================================

# adding a picture
document.add_picture('image-filename.png')
document.add_picture('image-filename.png', width=Inches(1.0))

# applying a paragraph style, bold , italic, and a character style
document.add_paragraph('Lorem ipsum dolor sit amet.', style='ListBullet')
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
paragraph.style = 'List Bullet'

paragraph = document.add_paragraph('Lorem ipsum ')
paragraph.add_run('dolor sit amet.')
paragraph = document.add_paragraph('Lorem ipsum ')
run = paragraph.add_run('dolor')
run.bold = True
paragraph.add_run(' sit amet.')
paragraph.add_run('dolor').bold = True

# is equivalent to:

run = paragraph.add_run('dolor')
run.bold = True

# except you don't have a reference to `run` afterward

paragraph = document.add_paragraph()
paragraph.add_run('Lorem ipsum ')
paragraph.add_run('dolor').bold = True
paragraph.add_run(' sit amet.')

paragraph = document.add_paragraph('Normal text, ')
paragraph.add_run('text with emphasis.', 'Emphasis')

paragraph = document.add_paragraph('Normal text, ')
run = paragraph.add_run('text with emphasis.')
run.style = 'Emphasis'